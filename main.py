# First here is the list of needed imports
import tkinter as tk
from tkinter import ttk # This import allows to use the tcl files to improve the visual of the GUI
import pandas as pd
import openpyxl # This allows to create a calendar for the new week in an Excel file
from datetime import datetime
import calendar
from docx import Document # This is needed to create the Word file
from morning import market_data_morning # This imports the morning function from a separate file
from news_AI import news_search # This imports the AI trained search engine
from news_AI import AI_report # This imports the pre-trained LLM that creates reports with the selected information

# TEST -> This values are just to try running the code for a given date and hour.
test_date = '2023-11-24'
test_time = '9'

# Store the current date, can compare with the day of publication given a certain format
c = datetime.now()
actual_date = c.strftime("%Y-%m-%d")

# Store the current time, so we can compare with the hour of publication of the rest of the variables.
current_time = c.strftime('%H')

#### The functions allow to change from one page to another:
# Function to switch to the "Create Calendar" page
def switch_to_calendario_page():
    """
    This function is used to switch to the Create Calendar page, it will be associated with an
    actionable click
    """
    frame1.pack_forget()  # Hide the current frame
    frame_calendario.pack()  # Show the "Create Calendar" page

# Function to switch to the "Create Report" page
def switch_to_reporte_page():
    """
    This function is used to switch to the Create Report page, it will be associated with an
    actionable click
    """
    frame1.pack_forget()  # Hide the current frame
    frame_reporte.pack()  # Show the "Create Report" page

# Function to go back to the main menu
def back_to_main_menu():
    """
    This function when associated with a click, allow to come back to the main menu.
    """
    frame_calendario.pack_forget()  # Hide the "Create Calendar" page
    frame_reporte.pack_forget()  # Hide the "Create Report" page
    frame1.pack()  # Show the main menu

#### Functions that allows to load and store data
def load_data():
    """
    Load the data in the excel file, and does this for the panel of data in the CREAR CALENDARIO page
    This function is useful just for one of the treeview that is used in this GUI.
    """
    path = "C:\\Users\\nesto\OneDrive\Documentos\pythonProject\Forest-ttk-theme\eco_event.xlsx"
    workbook = openpyxl.load_workbook(path)
    sheet =  workbook.active

    list_values = list(sheet.values)
    for col_name in list_values[0]:
        treeview.heading(col_name, text=col_name)

    for value_tuple in list_values[1:]:
        treeview.insert('',tk.END, values= value_tuple)

path = "C:\\Users\\nesto\OneDrive\Documentos\pythonProject\Forest-ttk-theme\eco_event.xlsx"
calendario_excel = pd.read_excel(path, header=0)
calendario_excel['Valor'] = 'Pendiente'
def load_data2(df):
    """
    Load the data in the excel file, and does this for the panel of data in the CREAR REPORTE page
    It is needed to close the program to read the new additions in the calendar
    This functions is useful for the second treeview used in this GUI
    """
    for i, row in df.iterrows():
        treeview2.insert('','end',values=tuple(row))

# function to insert a new row of data
def insert_row():
    """
    This function allow to insert the annotated data in a new row
    This function overwrite the file that save the economic events
    """
    event = event_entry.get()
    country = country_entry.get()
    date = date_entry.get()
    hour = hour_entry.get()
    prev = prev_entry.get()

    # insert row in excel sheet
    path = "C:\\Users\\nesto\OneDrive\Documentos\pythonProject\Forest-ttk-theme\eco_event.xlsx"
    workbook = openpyxl.load_workbook(path)
    sheet = workbook.active
    row_values = [event, country, date, hour, prev]
    sheet.append(row_values)
    workbook.save(path)

    # insert row in treeview
    treeview.insert('',tk.END,values=row_values)

    # Clear the values
    event_entry.set(events_list[0])
    country_entry.set(country_list[0])
    date_entry.delete(0,'end')
    date_entry.insert(0,'Fecha dd/m/aa')
    hour_entry.delete(0,'end')
    hour_entry.insert(0,'Hora')
    prev_entry.delete(0,'end')
    prev_entry.insert(0,'Previsión')

# Function to add the actual data to the calendar
i = 0

def insert_data():
    """
    Insert the actual value in the previously defined row
    """
    global i
    dato_public = data_entry.get()

    # insert data in a pandas df, iterate  loop by loop
    #calendario_excel.iloc[i,5] =dato_public
    filtered_calendario_excel.iloc[i,5] = dato_public


    # insert row in treeview
        # notice this is independent from the previous pandas df
    treeview3.insert('', tk.END,value=dato_public)
    i = i +1

    # clear the values:
    data_entry.delete(0,'end')
    data_entry.insert(0, 'Dato publicado')

# Function to create a excel calendar file
def excel_calendar():
    """
    This function is currently in development, the good looking Excel file calendar is not included yet
    """
    path = "C:\\Users\\nexecute_selected_functionesto\OneDrive\Documentos\pythonProject\Forest-ttk-theme\eco_event.xlsx"
    calendario_excel = pd.read_excel(path,header=0)
    #calendario_excel['Evento']
    #calendario_excel['País']
    calendario_excel['Fecha'] = pd.to_datetime(calendario_excel['Fecha'],format='%d/%m/%y')
    #calendario_excel['Hora']
    #calendario_excel['Previsión']

# This is the function that allows to change the time in the same day, morning, mid, end season.
def execute_selected_function():
    """
    This function allow to change easily the settings, depending on the time: morning, mid or closure.
    Called with the Publish call
    """
    selected_option = time_entry.get()
    if selected_option == "Market opening":
        #result_label.config(text="Function for Option 1 executed")
        market_data_morning()
        # Add your code for Option 1 here
    elif selected_option == "Half market season":
        print('not yet')
        # result_label.config(text="Function for Option 2 executed")
        # Add your code for Option 2 here
    elif selected_option == "Market closure":
        print('not yet')
        # result_label.config(text="Function for Option 3 executed")
        # Add your code for Option 3 here





#### List of countries that are allowed, and of the selected items, markets and hours
country_list = ['Select a Country','Spain','Italy','Germany','French']
events_list = ['Select an Event','CPI','HCPI','CPI core','IPI','GPD','Unemployment %','Current Account Balance','Trade Balance','PMI manufacturer','PMI services','PMI complex','Industrial Production','Retail Sales','Consumer Confidence','Business Confidence Index']
market_list = ['Select a Market','Ibex 35','U.S.A.', 'Asia'] # Currently just the Ibex market is valid
time_list = ['Select a time','Market opening','Half market season','Market closure']

#### Create the main window
root = tk.Tk()

# Add the style modification to the general frame, if prefer none of other change this, or comment this section
style = ttk.Style(root) # style variable whose parent is the root itself
root.tk.call('source','forest-light.tcl')
root.tk.call('source','forest-dark.tcl')
style.theme_use('forest-dark')

# Title of the whole program
root.title("Autonomous Report Tool")

# Main menu frame
frame1 = ttk.Frame(root)
frame1.pack() # make the frame resizable

# Create the widget Frame that stores both button
widget_frame1 = ttk.LabelFrame(frame1,text='Select an action: ')
widget_frame1.grid(row=0,column=0, padx=20, pady=10)

# Create the "Create Calendar" button on the main menu
button_crear_calendario = ttk.Button(widget_frame1, text="Create a Calendar", command=switch_to_calendario_page)
#button_crear_calendario.pack(side=tk.LEFT)
button_crear_calendario.grid(row=0, column=0,padx=5,pady=[0,3],sticky= 'ew')

# Create the "Create Report" button on the main menu
button_crear_reporte = ttk.Button(widget_frame1, text="Create a Report", command=switch_to_reporte_page)
#button_crear_reporte.pack(side=tk.RIGHT)
button_crear_reporte.grid(row=0, column=1,padx=5,pady=[0,3],sticky='ew')



#### "Create Calendar" page
frame_calendario = ttk.Frame(root)
#label_calendario = ttk.Label(frame_calendario, text="Página de Crear Calendario")
label_calendario = ttk.Label(frame_calendario)
label_calendario.pack()
label_calendario_text = ttk.LabelFrame(label_calendario,text='Please, indicate the next week economic events: ')
label_calendario_text.grid(row=0,column=0, padx=20, pady=10)

# Entry line to add Economics Events
event_entry = ttk.Combobox(label_calendario_text, values=events_list)
event_entry.current(0)
event_entry.grid(row=0,column=0, padx=[5,3],pady=[0,3], sticky='ew')

# Entry line to add country
country_entry = ttk.Combobox(label_calendario_text, values=country_list)
country_entry.current(0)
country_entry.grid(row=0,column=1,padx=[3,3],pady=[0,3], sticky='ew')

# Entry line to add date
date_entry = ttk.Entry(label_calendario_text)
date_entry.insert(0,'Date dd/mm/aa')
date_entry.bind("<FocusIn>", lambda e:date_entry.delete('0','end'))
date_entry.grid(row=0, column=2,padx=[3,3],pady=[0,3], sticky='ew')

# Entry line to add hour
hour_entry = ttk.Entry(label_calendario_text)
hour_entry.insert(0,'Hour')
hour_entry.bind("<FocusIn>", lambda e:hour_entry.delete('0','end'))
hour_entry.grid(row=0, column =3,padx=[3,3],pady=[0,3], sticky='ew')

# Entry the estimated value:
prev_entry = ttk.Entry(label_calendario_text)
prev_entry.insert(0,'Forecast')
prev_entry.bind("<FocusIn>", lambda e:prev_entry.delete('0','end'))
prev_entry.grid(row=0, column = 4,padx=[3,5],pady=[0,3], sticky='ew')

# Insert data
insert = ttk.Button(label_calendario_text, text='Add', command=insert_row)
insert.grid(row=1, column=1, pady=[3,5], sticky='nsew')

# Create excel calendar (again, this functionality is not included yet)
create = ttk.Button(label_calendario_text, text='Create',style='Accent.TButton', command=excel_calendar)
create.grid(row=1,column=3, pady=[3,5], sticky='nsew')

# visualize the inserted data
treeFrame = ttk.Frame(label_calendario)
treeFrame.grid(row=1,column=0,pady=10)
treeScroll = ttk.Scrollbar(treeFrame)
treeScroll.pack(side='right',fill='y')

columns = ('Evento','País','Fecha','Hora','Previsión')
treeview = ttk.Treeview(treeFrame,show='headings',
                        yscrollcommand=treeScroll.set,columns=columns, height=13)
treeview.column('Evento',width=200)
treeview.column('País', width=100)
treeview.column('Fecha', width=100)
treeview.column('Hora', width=50)
treeview.column('Previsión', width=70)
treeview.pack()
treeScroll.config(command=treeview.yview)
load_data()

# Go to main menu
back_button_calendario = tk.Button(frame_calendario, text="Go back to main menu", command=back_to_main_menu)
back_button_calendario.pack()


#### "Create Report" page
frame_reporte = ttk.Frame(root)
label_reporte = ttk.Label(frame_reporte)
label_reporte.pack()
label_reporte_text = ttk.LabelFrame(label_reporte,text='Please, indicate the published data: ')
label_reporte_text.grid(row=0,column=0, padx=20, pady=10)

calendario_excel['Fecha2'] = pd.to_datetime(calendario_excel['Fecha'], format='%d/%m/%Y')
filtered_calendario_excel = calendario_excel[calendario_excel['Fecha2'] == test_date]

# show past generated  data:
treeFrame2 = ttk.Frame(label_reporte)
treeFrame2.grid(row=1,column=0,padx= 20, pady=10)
columns2 = ('Evento','País','Fecha','Hora','Previsión')
treeview2 = ttk.Treeview(treeFrame2,show='headings',
                        yscrollcommand=treeScroll.set,columns=columns2, height=10)
treeview2.heading('Evento',text='Event')
treeview2.column('Evento',width=200)
treeview2.heading('País',text= 'Country')
treeview2.column('País', width=100)
treeview2.heading('Fecha',text='Date')
treeview2.column('Fecha', width=100)
treeview2.heading('Hora',text= 'Hour')
treeview2.column('Hora', width=50)
treeview2.heading('Previsión',text='Forecast')
treeview2.column('Previsión', width=70)
treeview2.pack(side='left',fill='both',expand=True)
load_data2(filtered_calendario_excel)

# Show the inserted data, when the actual value comes
treeview3 = ttk.Treeview(treeFrame2, show='headings',yscrollcommand=treeScroll.set, columns=('Valor'), height=10)
treeview3.heading('Valor', text = 'Valor')
treeview3.column('Valor',width=15)
treeview3.pack(side='left',padx=[5,0])

# Entry actual data:
data_entry = ttk.Entry(label_reporte_text,width=30)
data_entry.insert(0,'Published data')
data_entry.bind("<FocusIn>", lambda e:data_entry.delete('0','end'))
data_entry.grid(row=1, column=0, padx=[100,100], pady=[6,3], sticky='ew')

# Entry line to add country:
market_entry = ttk.Combobox(label_reporte_text, values=market_list)
market_entry.current(0)
market_entry.grid(row=3,column=0,padx=[70,500],pady=[5,10], sticky='ew')

# Entry line to add the time in the day and in the season:
time_entry = ttk.Combobox(label_reporte_text, values = time_list)
time_entry.current(0)
time_entry.grid(row=3, column=0, padx=[290,250], pady =[5,10], sticky='ew')

# Button to enter data
insert = ttk.Button(label_reporte_text, text='Add', command= insert_data)
insert.grid(row=2, column=0,pady=5)


# Button to write the report itself once the data has been stored
#publish = ttk.Button(label_reporte_text, text='Publicar', style='Accent.TButton', command= market_data_morning)
publish = ttk.Button(label_reporte_text, text='Publish', style='Accent.TButton', command= execute_selected_function)
publish.grid(row=3,column = 0, padx=[500,0],pady=[5,10])

# Go to main menu
back_button_reporte = tk.Button(frame_reporte, text="Go back to main menu", command=back_to_main_menu)
back_button_reporte.pack()

# Start the GUI main loop
root.mainloop()

# here comes the new financial news information
descripciones = news_search()
financial_news = AI_report(descripciones)

# Obtain the news from the AI search engine
a_full, b_full = market_data_morning()

document = Document()
document.add_paragraph(financial_news)
document.add_paragraph(a_full)
#document.add_paragraph(macro_indicators)
document.add_paragraph(b_full)
document.save('morning_report.docx')
